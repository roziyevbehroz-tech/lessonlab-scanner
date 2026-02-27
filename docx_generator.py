from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_test_docx(title, questions, output_path):
    """Test savollarini .docx formatida yaratadi."""
    doc = Document()

    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    info.add_run(f"Savollar soni: {len(questions)} ta").italic = True
    info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph("_" * 50)

    for i, q in enumerate(questions, 1):
        doc.add_paragraph().add_run(f"{i}. {q['text']}").bold = True
        for opt in q['options']:
            run = doc.add_paragraph(style='List Bullet').add_run(opt['text'])
            if opt['is_correct']:
                run.font.color.rgb = RGBColor(0, 150, 0)
                run.bold = True

    doc.save(output_path)
    return output_path
